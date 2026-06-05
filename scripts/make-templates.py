# Gera templates .docx (com campos {{...}}) a partir dos contratos-amostra.
# Uso: py -3 scripts/make-templates.py
# Substitui os valores das amostras por placeholders e verifica que nenhum
# dado pessoal da amostra sobrou no template.
import re
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC = Path(r"C:\Users\hsgus\OneDrive\Downloads Web\teste_docs\contratos")
OUT = Path(r"C:\Users\hsgus\OneDrive\Claude Code\dashboard_Auri\auri-dashboard\public\contratos-modelos")
OUT.mkdir(parents=True, exist_ok=True)

def apply_to_paragraph(p, repls):
    full = "".join(r.text for r in p.runs)
    if not full.strip():
        return
    new = full
    for pat, rep in repls:
        new = pat.sub(rep, new)
    if new != full and p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""

def _tables_paragraphs(tables):
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    yield from _tables_paragraphs(doc.tables)
    for sec in doc.sections:
        for hf in (sec.header, sec.footer, sec.first_page_header,
                   sec.first_page_footer, sec.even_page_header, sec.even_page_footer):
            for p in hf.paragraphs:
                yield p
            yield from _tables_paragraphs(hf.tables)
    # parágrafos dentro de caixas de texto (textboxes) do corpo e de cabeçalhos
    parts = [doc.element.body]
    for sec in doc.sections:
        for hf in (sec.header, sec.footer, sec.first_page_header,
                   sec.first_page_footer, sec.even_page_header, sec.even_page_footer):
            parts.append(hf._element)
    for part in parts:
        for tx in part.iter(qn("w:txbxContent")):
            for p_el in tx.iter(qn("w:p")):
                yield Paragraph(p_el, None)

def process(src_name, out_name, repls, leftovers):
    doc = Document(SRC / src_name)
    for p in iter_paragraphs(doc):
        apply_to_paragraph(p, repls)
    doc.save(OUT / out_name)
    full = "\n".join(p.text for p in iter_paragraphs(doc))
    resto = [tok for tok in leftovers if tok.lower() in full.lower()]
    tags = sorted(set(re.findall(r"{{[a-z_]+}}", full)))
    print(f"\n== {out_name} ==")
    print("  placeholders:", ", ".join(tags) if tags else "(NENHUM!)")
    if resto:
        print("  !! SOBROU dado da amostra:", resto)
    else:
        print("  OK — nenhum dado da amostra restante")
    return not resto and bool(tags)

repls_pf = [
    (re.compile(r"Avenida 5,.*?CEP:\s*75\.832-011\.?"), "{{endereco_completo}}"),
    (re.compile(r"Avenida 5,.*?Mineiros, Goiás"), "{{endereco_completo}}"),
    (re.compile(r"Mineiros,\s*\d{1,2} de [A-Za-zçÇ]+ de \d{4}"), "{{local_data}}"),
    (re.compile(r"034\.060\.168-04"), "{{titular_cpf_cnpj}}"),
    (re.compile(r"Edson de Clemente"), "{{titular_nome_ou_razao}}"),
    (re.compile(r"\(64\)\s*99605-7673"), "{{titular_telefone}}"),
    (re.compile(r"renatadeclemente@hotmail\.com"), "{{titular_email}}"),
    (re.compile(r"1390102730"), "{{uc}}"),
    (re.compile(r"\b1000\b"), "{{consumo_medio_kwh}}"),
    (re.compile(r"AE29"), "{{numero_contrato}}"),
]
leftovers_pf = ["034.060.168-04", "Edson", "99605-7673", "renatadeclemente", "1390102730", "Avenida 5", "75.832-011", "AE29"]

repls_pj = [
    (re.compile(r"Av\. 05,.*?Mineiros, Goiás"), "{{endereco_completo}}"),
    (re.compile(r"Mineiros,\s*\d{1,2} de [A-Za-zçÇ]+ de \d{4}"), "{{local_data}}"),
    (re.compile(r"57\.893\.714/0001-93"), "{{titular_cpf_cnpj}}"),
    (re.compile(r"Condomínio Residencial Ararauna", re.IGNORECASE), "{{titular_nome_ou_razao}}"),
    (re.compile(r"\(64\)\s*99938-3041"), "{{titular_telefone}}"),
    (re.compile(r"1390058004"), "{{uc}}"),
    (re.compile(r"\b2100\b"), "{{consumo_medio_kwh}}"),
    (re.compile(r"AE34"), "{{numero_contrato}}"),
]
leftovers_pj = ["57.893.714/0001-93", "Ararauna", "99938-3041", "1390058004", "Av. 05", "75.830-001", "AE34"]

# ── Locação de Equipamento PF ──────────────────────────────────────────────────
repls_pf_locacao = [
    (re.compile(r"AE29"), "{{numero_contrato}}"),
    (re.compile(r"Edson de Clemente"), "{{titular_nome_ou_razao}}"),
    (re.compile(r"Brasileiro\s+\(a\),\s*Casado,\s*Agropecuarista"), "{{titular_nacionalidade}}, {{titular_estado_civil}}, {{titular_profissao}}"),
    (re.compile(r"034\.060\.168-04"), "{{titular_cpf_cnpj}}"),
    (re.compile(r"4915774"), "{{titular_rg}}"),
    (re.compile(r"SSP\s+SP"), "{{titular_rg_orgao}}"),
    (re.compile(r"Avenida\s+5,\s*Qd\s*10\s*-\s*Lt\s*17,\s*Setor Mundinho,\s*Mineiros\s*-\s*GO,\s*CEP:\s*75\.832-011"), "{{endereco_completo}}"),
    (re.compile(r"ricardo\.declemente@gmail\.com"), "{{titular_email}}"),
    (re.compile(r"renatadeclemente@hotmail\.com@gmail\.com"), "{{titular_email}}"),
    (re.compile(r"renatadeclemente@hotmail\.com"), "{{titular_email}}"),
    (re.compile(r"\(64\)\s*99605-7673"), "{{titular_telefone}}"),
    (re.compile(r"\b1390102730\b"), "{{uc}}"),
    (re.compile(r"DESC_GAR é igual a 10%"), "DESC_GAR é igual a {{desconto_garantido_pct}}%"),
    (re.compile(r"Energia contratada \(MWh/ano\)"), "Energia contratada (kWh/ano)"),
    (re.compile(r"\b12\.000\b"), "{{energia_contratada_kwh_ano}}"),
    (re.compile(r"Mineiros,\s*22 de Maio de 2025"), "{{local_data}}"),
]
leftovers_pf_locacao = ["034.060.168-04", "Edson de Clemente", "99605-7673", "renatadeclemente", "ricardo.declemente", "1390102730", "AE29", "4915774"]

# ── Locação de Equipamento PJ ──────────────────────────────────────────────────
repls_pj_locacao = [
    (re.compile(r"AE34"), "{{numero_contrato}}"),
    (re.compile(r"CONDOMINIO RESIDENCIAL ARARAUNA"), "{{titular_nome_ou_razao}}"),
    (re.compile(r"Condomínio Residencial Ararauna"), "{{titular_nome_ou_razao}}"),
    (re.compile(r"57\.893\.714/0001-93"), "{{titular_cpf_cnpj}}"),
    (re.compile(r"Av\. 05,\s*Q\.9,\s*L\.15,\s*Setor Martins,\s*CEP:\s*75\.830-001,\s*Mineiros,\s*Goiás\s*,?\s*representado"), "{{endereco_completo}}, representado"),
    (re.compile(r"Av\. 05,\s*Q\.9,\s*L\.15,\s*Setor Martins,\s*CEP:\s*75\.830-001,\s*Mineiros,\s*Goiás"), "{{endereco_completo}}"),
    (re.compile(r"Av\. 05,\s*Q\.9,\s*L\.15,\s*Residencial Ararauna,\s*Setor Martins,\s*CEP:\s*75\.830-001,\s*Mineiros,\s*Goiás"), "{{endereco_completo}}"),
    (re.compile(r"por seu síndico,"), "por seu {{rep_cargo}},"),
    (re.compile(r"Thais Cristina Sousa Fernandes"), "{{rep_nome}}"),
    (re.compile(r"Brasileira,\s*Casada,\s*Engenheira Florestal"), "{{rep_qualificacao}}"),
    (re.compile(r"009\.814\.481-27"), "{{rep_cpf}}"),
    (re.compile(r"4241657\s*/\s*PC GO"), "{{rep_rg}} / {{rep_rg_orgao}}"),
    (re.compile(r"Rua Antônio Rodrigues da Silva,\s*Q\.8,\s*L\.5,\s*AP\.903,\s*Setor Rodrigues,\s*Mineiros,\s*GO,\s*CEP\s*75\.832-127"), "{{rep_endereco}}"),
    (re.compile(r"\b1390058004\b"), "{{uc}}"),
    (re.compile(r"\(64\)\s*99938-3041"), "{{titular_telefone}}"),
    (re.compile(r"DESC_GAR é igual a 10%"), "DESC_GAR é igual a {{desconto_garantido_pct}}%"),
    (re.compile(r"\b25\.200\b"), "{{energia_contratada_kwh_ano}}"),
    (re.compile(r"Mineiros,\s*03 de Junho de 2025"), "{{local_data}}"),
]
leftovers_pj_locacao = ["57.893.714/0001-93", "Ararauna", "009.814.481-27", "4241657", "Thais Cristina", "99938-3041", "1390058004", "AE34", "25.200"]

# ── Locação de Imóvel PF ───────────────────────────────────────────────────────
repls_pf_aluguel = [
    (re.compile(r"Edson de Clemente"), "{{titular_nome_ou_razao}}"),
    (re.compile(r"Brasileiro\s+\(a\),\s*Casado,\s*Agropecuarista"), "{{titular_nacionalidade}}, {{titular_estado_civil}}, {{titular_profissao}}"),
    (re.compile(r"034\.060\.168-04"), "{{titular_cpf_cnpj}}"),
    (re.compile(r"4915774"), "{{titular_rg}}"),
    (re.compile(r"SSP\s+GO"), "{{titular_rg_orgao}}"),
    (re.compile(r"Rua 04,\s*Qd\s*8\s*-\s*Lt\s*04,\s*Setor Mundinho,\s*Mineiros\s*-\s*GO,\s*CEP:\s*75\.832-011"), "{{endereco_completo}}"),
    (re.compile(r"renatadeclemente@hotmail\.com"), "{{titular_email}}"),
    # endereço do imóvel (cláusula 1)
    (re.compile(r"Avenida\s+5,\s*Qd\s*10\s*-\s*Lt\s*17,\s*Setor Mundinho,\s*Mineiros\s*-\s*GO,\s*CEP:\s*75\.832-011"), "{{aluguel_imovel_endereco}}"),
    # prazo numérico (cláusula 2: "é de 12 meses")
    (re.compile(r"é de 12 meses"), "é de {{aluguel_prazo_meses}} meses"),
    # prazo por extenso (cláusula 3: "doze meses de locação")
    (re.compile(r"doze meses de locação"), "{{aluguel_prazo_extenso}} meses de locação"),
    # data de início
    (re.compile(r"23/09/2024"), "{{aluguel_data_inicio}}"),
    # valor total (cláusula 3)
    (re.compile(r"R\$\s*22\.000,00\s*\(vinte e dois mil reais\)"), "R$ {{aluguel_valor_total_fmt}} ({{aluguel_valor_total_extenso}})"),
    # valor mensal (aparece 2×)
    (re.compile(r"R\$\s*1\.000,00\s*\(um mil reais\)"), "R$ {{aluguel_valor_mensal_fmt}} ({{aluguel_valor_mensal_extenso}})"),
    (re.compile(r"Goiânia,\s*22 de Maio de 2025"), "{{local_data}}"),
]
leftovers_pf_aluguel = ["034.060.168-04", "Edson de Clemente", "renatadeclemente", "4915774", "22.000", "1.000,00", "23/09/2024", "Rua 04, Qd 8"]

# ── Locação de Imóvel PJ ───────────────────────────────────────────────────────
repls_pj_aluguel = [
    (re.compile(r"Condomínio Residencial Ararauna"), "{{titular_nome_ou_razao}}"),
    (re.compile(r"57\.893\.714/0001-93"), "{{titular_cpf_cnpj}}"),
    (re.compile(r"Av\. 05,\s*Q\.9,\s*L\.15,\s*Setor Martins,\s*CEP:\s*75\.830-001,\s*Mineiros,\s*Goiás"), "{{endereco_completo}}"),
    (re.compile(r"por seu síndico,"), "por seu {{rep_cargo}},"),
    (re.compile(r"Thais Cristina Sousa Fernandes"), "{{rep_nome}}"),
    (re.compile(r"Brasileira,\s*Casada,\s*Síndica"), "{{rep_qualificacao}}"),
    (re.compile(r"009\.814\.481-27"), "{{rep_cpf}}"),
    (re.compile(r"4241657\s*/\s*PC GO"), "{{rep_rg}} / {{rep_rg_orgao}}"),
    (re.compile(r"Rua Antônio Rodrigues da Silva,\s*Q\.8,\s*L\.5,\s*Res\.\s*Mirante do Cerrado"), "{{rep_endereco}}"),
    # endereço do imóvel (cláusula 1) — tem "Residencial Ararauna" explícito
    (re.compile(r"Av\. 05,\s*Q\.9,\s*L\.15,\s*Residencial Ararauna,\s*Setor Martins,\s*CEP:\s*75\.830-001,\s*Mineiros,\s*Goiás"), "{{aluguel_imovel_endereco}}"),
    # prazo numérico (cláusula 2)
    (re.compile(r"é de 120 meses"), "é de {{aluguel_prazo_meses}} meses"),
    # prazo por extenso (cláusula 3)
    (re.compile(r"cento e vinte meses de locação"), "{{aluguel_prazo_extenso}} meses de locação"),
    # data de início
    (re.compile(r"04/06/2025"), "{{aluguel_data_inicio}}"),
    # valor total (cláusula 3)
    (re.compile(r"R\$\s*120\.000,00\s*\(cento e vinte mil reais\)"), "R$ {{aluguel_valor_total_fmt}} ({{aluguel_valor_total_extenso}})"),
    # valor mensal (aparece 2×)
    (re.compile(r"R\$\s*1\.000,00\s*\(um mil reais\)"), "R$ {{aluguel_valor_mensal_fmt}} ({{aluguel_valor_mensal_extenso}})"),
    (re.compile(r"Mineiros,\s*0?3 de junho de 2025"), "{{local_data}}"),
]
leftovers_pj_aluguel = ["57.893.714/0001-93", "Ararauna", "009.814.481-27", "4241657", "Thais Cristina", "1.000,00", "120.000", "04/06/2025"]

ok1 = process("pf_adesao.docx",   "adesao-pf.docx",       repls_pf,         leftovers_pf)
ok2 = process("pj_adesao.docx",   "adesao-pj.docx",       repls_pj,         leftovers_pj)
ok3 = process("pf_locacao.docx",  "locacao-equip-pf.docx", repls_pf_locacao, leftovers_pf_locacao)
ok4 = process("pj_locacao.docx",  "locacao-equip-pj.docx", repls_pj_locacao, leftovers_pj_locacao)
ok5 = process("pf_aluguel.docx",  "aluguel-pf.docx",       repls_pf_aluguel, leftovers_pf_aluguel)
ok6 = process("pj_aluguel.docx",  "aluguel-pj.docx",       repls_pj_aluguel, leftovers_pj_aluguel)
sys.exit(0 if all([ok1, ok2, ok3, ok4, ok5, ok6]) else 1)
